from telegram import Update

from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    MessageHandler,
    ContextTypes,
    filters
)

import json

import time 

import psutil

from deep_translator import GoogleTranslator

import langid

import re

from graph import app_graph

# =========================================
# NEW IMPORTS FOR FILE UPLOAD
# =========================================

import os

from langchain_text_splitters import RecursiveCharacterTextSplitter

from langchain_core.documents import Document

from langchain_huggingface import HuggingFaceEmbeddings

from langchain_community.vectorstores import Chroma

from pypdf import PdfReader

from docx import Document as DocxDocument


# =========================================
# TELEGRAM BOT TOKEN
# =========================================

BOT_TOKEN = "8957560769:AAHDHPgJ290Ht1RW_OMepf2sHgrsA9I75vY"


# =========================================
# PATH SETUP
# =========================================

BASE_DIR = os.path.dirname(
    os.path.abspath(__file__)
)

UPLOAD_DIR = os.path.join(
    BASE_DIR,
    "uploads"
)

VECTOR_DB_DIR = os.path.join(
    BASE_DIR,
    "vector_db"
)

os.makedirs(UPLOAD_DIR, exist_ok=True)


# =========================================
# EMBEDDINGS + VECTOR DB
# =========================================

embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)

db = Chroma(
    persist_directory=VECTOR_DB_DIR,
    embedding_function=embeddings
)

def is_english_text(text):

    return bool(
        re.fullmatch(r"[A-Za-z0-9\s.,!?'-]+", text)
    )



# =========================================
# STATS FUNCTIONS
# =========================================

STATS_FILE = "stats.json"

def load_stats():

    with open(STATS_FILE, "r") as f:

        return json.load(f)

def save_stats(stats):

    with open(STATS_FILE, "w") as f:

        json.dump(stats, f, indent=4)

# =========================================
# START COMMAND
# =========================================

async def start(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE
):

    await update.message.reply_text(
        "Hello! I am your SSN College AI Assistant."
    )


# =========================================
# HANDLE TEXT MESSAGE
# =========================================

async def handle_message(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE
):

    user_message = update.message.text

    print("\n========================")
    print("Original Message:")
    print(user_message)

    # =====================================
    # LANGUAGE DETECTION
    # =====================================

    try:

    # If pure English characters
        if is_english_text(user_message):

            detected_lang = "en"
            confidence = 1

        else:

            detected_lang, confidence = langid.classify(
                user_message
            )

    except:

        detected_lang = "en"

    print("Detected Language:", detected_lang)
    print("Confidence:", confidence)

    print("Detected Language:", detected_lang)

    # =====================================
    # TRANSLATE TO ENGLISH
    # =====================================

    if detected_lang != "en":

        translated_question = GoogleTranslator(
            source="auto",
            target="en"
        ).translate(user_message)

    else:

        translated_question = user_message

    print("Translated Question:")
    print(translated_question)

    # =====================================
    # CHATBOT
    # =====================================
    start_time = time.time()

    result = app_graph.invoke({
        "question": translated_question
    })

    bot_reply = result["answer"]

    end_time = time.time()

    latency = end_time - start_time

    # Approx token calculation
    token_count = len(bot_reply.split())

    tokens_per_sec = token_count / latency

    # RAM usage
    ram = psutil.virtual_memory().used / (1024 ** 3)

    stats = load_stats()

    stats["total_queries"] += 1

    stats["last_question"] = user_message

    stats["last_language"] = detected_lang

    stats["tokens_per_second"] = round(tokens_per_sec, 2)

    stats["ram_usage"] = round(ram, 2)

    count = stats["total_queries"]

    old_avg = stats["average_latency"]

    new_avg = ((old_avg * (count - 1)) + latency) / count

    stats["average_latency"] = round(new_avg, 2)

    save_stats(stats)

    print(f"Latency: {latency:.2f} sec")

    print(f"Tokens/sec: {tokens_per_sec:.2f}")

    print(f"RAM Usage: {ram:.2f} GB")


    print("Bot Reply:")
    print(bot_reply)

    # =====================================
    # TRANSLATE BACK
    # =====================================

    if detected_lang != "en":

        final_reply = GoogleTranslator(
            source="en",
            target=detected_lang
        ).translate(bot_reply)

    else:

        final_reply = bot_reply

    print("Final Reply:")
    print(final_reply)

    await update.message.reply_text(final_reply)


# =========================================
# HANDLE DOCUMENT UPLOAD
# =========================================

async def handle_document(
    update: Update,
    context: ContextTypes.DEFAULT_TYPE
):

    try:

        telegram_file = update.message.document

        filename = telegram_file.file_name

        print(f"\nReceiving File: {filename}")

        # =================================
        # DOWNLOAD FILE
        # =================================

        file = await context.bot.get_file(
            telegram_file.file_id
        )

        save_path = os.path.join(
            UPLOAD_DIR,
            filename
        )

        await file.download_to_drive(save_path)

        print(f"Saved: {save_path}")

        # =================================
        # EXTRACT TEXT
        # =================================

        text = ""

        # TXT
        if filename.endswith(".txt"):

            with open(
                save_path,
                "r",
                encoding="utf-8"
            ) as f:

                text = f.read()

        # PDF
        elif filename.endswith(".pdf"):

            pdf = PdfReader(save_path)

            for page in pdf.pages:

                extracted = page.extract_text()

                if extracted:

                    text += extracted + "\n"

        # DOCX
        elif filename.endswith(".docx"):

            doc = DocxDocument(save_path)

            for para in doc.paragraphs:

                text += para.text + "\n"

        else:

            await update.message.reply_text(
                "Unsupported file type."
            )

            return

        # =================================
        # VALIDATE
        # =================================

        if len(text.strip()) < 20:

            await update.message.reply_text(
                "Document is empty."
            )

            return

        # =================================
        # FAQ CHUNKING
        # =================================

        chunks = text.split("\n\n")

        documents = [

            Document(page_content=chunk.strip())

            for chunk in chunks

            if chunk.strip()
        ]

        # =================================
        # ADD TO VECTOR DB
        # =================================

        db.add_documents(documents)


        stats = load_stats()

        stats["documents_uploaded"] += 1

        stats["total_chunks"] += len(documents)

        save_stats(stats)


        print("Documents added to vector DB")

        await update.message.reply_text(
            f"{filename} uploaded successfully.\n"
            f"Chunks added: {len(documents)}"
        )

    except Exception as e:

        print("Upload Error:", e)

        await update.message.reply_text(
            "File upload failed."
        )


# =========================================
# MAIN
# =========================================

def main():

    app = ApplicationBuilder().token(
        BOT_TOKEN
    ).build()

    # Start command
    app.add_handler(
        CommandHandler("start", start)
    )

    # Text messages
    app.add_handler(
        MessageHandler(
            filters.TEXT & ~filters.COMMAND,
            handle_message
        )
    )

    # =====================================
    # DOCUMENT HANDLER
    # =====================================

    app.add_handler(
        MessageHandler(
            filters.Document.ALL,
            handle_document
        )
    )

    print("Telegram bot running...")

    app.run_polling()


# =========================================
# RUN
# =========================================

if __name__ == "__main__":

    main()