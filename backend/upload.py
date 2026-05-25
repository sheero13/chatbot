import os

from fastapi import APIRouter, UploadFile, File

from langchain_text_splitters import RecursiveCharacterTextSplitter

from langchain_core.documents import Document

from langchain_huggingface import HuggingFaceEmbeddings

from langchain_community.vectorstores import Chroma

from pypdf import PdfReader

from docx import Document as DocxDocument

# =========================================
# ROUTER
# =========================================

router = APIRouter()

# Create uploads folder
os.makedirs("uploads", exist_ok=True)

# =========================================
# EMBEDDINGS
# =========================================

embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)

# =========================================
# LOAD VECTOR DB
# =========================================

db = Chroma(
    persist_directory="vector_db",
    embedding_function=embeddings
)

# =========================================
# UPLOAD ENDPOINT
# =========================================

@router.post("/upload")

async def upload_file(file: UploadFile = File(...)):

    filename = file.filename

    save_path = f"uploads/{filename}"

    # Save file
    with open(save_path, "wb") as f:

        f.write(await file.read())

    text = ""

    # TXT
    if filename.endswith(".txt"):

        with open(save_path, "r", encoding="utf-8") as f:

            text = f.read()

    # PDF
    elif filename.endswith(".pdf"):

        pdf = PdfReader(save_path)

        for page in pdf.pages:

            extracted = page.extract_text()

            if extracted:
                text += extracted + "\n"

    # DOCX
    elif filename.endswith(".docx"):

        doc = DocxDocument(save_path)

        for para in doc.paragraphs:

            text += para.text + "\n"

    else:

        return {
            "message": "Unsupported file type"
        }

    # Validate
    if len(text.strip()) < 20:

        return {
            "message": "Document is empty"
        }

    # Split
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )

    chunks = splitter.split_text(text)

    documents = [

        Document(page_content=chunk)

        for chunk in chunks
    ]

    # Add to vector DB
    db.add_documents(documents)

    return {
        "message": f"{filename} uploaded successfully",
        "chunks_added": len(documents)
    }